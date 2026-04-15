#!/usr/bin/env python3
"""
Build a Microsoft Word (.docx) report for one country: all nomination contestants
in that country on the active season, ranked like production migration (stars →
shares → likes → comments → views → contestant id). Marks Top 5 and who would
migrate to the next stage (default: top 5 per country → regional).

Read-only.

Usage (venv + PYTHONPATH same as other scripts):
  python3 scripts/export_country_nomination_word.py --round-id 3 --country Tanzania \\
      -o ~/reports/tanzania_nomination_round3.docx

  Single contest only:
  python3 scripts/export_country_nomination_word.py --round-id 3 --country Tanzania \\
      --contest-id 4 -o beauty_tz.docx
"""
from __future__ import annotations

import argparse
import sys
from typing import Dict, List, Optional, Set, Tuple

from sqlalchemy import and_, func, select

from app.models.contest import Contest
from app.models.contests import ContestSeason, ContestSeasonLink, Contestant, ContestantSeason, SeasonLevel
from app.models.round import Round, round_contests
from app.models.voting import ContestantVoting
from app.services.season_migration import SeasonMigrationService

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
except ImportError as e:  # pragma: no cover
    sys.exit(
        "python-docx is required. Install with: pip install python-docx\n"
        f"Import error: {e}"
    )


def _country_variants(label: str) -> Set[str]:
    """Match migration service country_filter aliases (lowercase)."""
    raw = label.strip().lower()
    variants = {raw}
    alias_map = {
        "tanzania": "tz",
        "tz": "tanzania",
        "uganda": "ug",
        "ug": "uganda",
        "kenya": "ke",
        "ke": "kenya",
    }
    if raw in alias_map:
        variants.add(alias_map[raw])
    return variants


def _points_map(
    db,
    season_id: int,
    contestant_ids: List[int],
    contest_id: int,
) -> Dict[int, int]:
    if not contestant_ids:
        return {}
    q = (
        db.query(
            ContestantVoting.contestant_id,
            func.coalesce(func.sum(ContestantVoting.points), 0).label("total_points"),
        )
        .filter(
            and_(
                ContestantVoting.season_id == season_id,
                ContestantVoting.contest_id == contest_id,
                ContestantVoting.contestant_id.in_(contestant_ids),
            )
        )
        .group_by(ContestantVoting.contestant_id)
    )
    rows = q.all()
    if rows:
        return {r.contestant_id: int(r.total_points or 0) for r in rows}
    rows2 = (
        db.query(
            ContestantVoting.contestant_id,
            func.coalesce(func.sum(ContestantVoting.points), 0).label("total_points"),
        )
        .filter(
            and_(
                ContestantVoting.contest_id == contest_id,
                ContestantVoting.contestant_id.in_(contestant_ids),
            )
        )
        .group_by(ContestantVoting.contestant_id)
        .all()
    )
    return {r.contestant_id: int(r.total_points or 0) for r in rows2}


def _contestant_label(c: Contestant) -> str:
    if c.title and str(c.title).strip():
        return str(c.title).strip()
    u = c.user
    if u is None:
        return f"Contestant #{c.id}"
    return (u.full_name or u.username or u.email or f"User #{u.id}")[:120]


def _next_level(current: SeasonLevel) -> Optional[SeasonLevel]:
    order = [
        SeasonLevel.CITY,
        SeasonLevel.COUNTRY,
        SeasonLevel.REGIONAL,
        SeasonLevel.CONTINENT,
        SeasonLevel.GLOBAL,
    ]
    try:
        i = order.index(current)
    except ValueError:
        return None
    if i >= len(order) - 1:
        return None
    return order[i + 1]


def _rank_tanzania_pool(
    db,
    *,
    season_id: int,
    contest_id: int,
    country_variants: Set[str],
) -> List[Tuple[Contestant, Dict[str, int], int]]:
    """
    All qualified active contestants on season for this contest whose
    Contestant.country matches variants, sorted by migration rules.
    Returns list of (contestant, engagement_dict, stars).
    """
    q = (
        db.query(Contestant)
        .join(ContestantSeason, ContestantSeason.contestant_id == Contestant.id)
        .filter(
            and_(
                ContestantSeason.season_id == season_id,
                ContestantSeason.is_active == True,
                Contestant.is_active == True,
                Contestant.is_deleted == False,
                Contestant.is_qualified == True,
                Contestant.season_id == contest_id,
            )
        )
    )
    all_c = q.all()
    pool = []
    for c in all_c:
        cc = (c.country or "").strip().lower()
        if cc and cc in country_variants:
            pool.append(c)
    ids = [c.id for c in pool]
    if not ids:
        return []
    pts = _points_map(db, season_id, ids, contest_id)
    eng = SeasonMigrationService._engagement_by_contestant(db, ids)
    ranked = sorted(
        pool,
        key=lambda c: (
            pts.get(c.id, 0),
            eng.get(c.id, {}).get("shares", 0),
            eng.get(c.id, {}).get("likes", 0),
            eng.get(c.id, {}).get("comments", 0),
            eng.get(c.id, {}).get("views", 0),
            -(c.id or 0),
        ),
        reverse=True,
    )
    out: List[Tuple[Contestant, Dict[str, int], int]] = []
    for c in ranked:
        e = eng.get(c.id, {})
        out.append((c, e, int(pts.get(c.id, 0))))
    return out


def _header_cell(cell, text: str) -> None:
    cell.text = text
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(10)


def _body_cell(cell, text: str, bold: bool = False) -> None:
    cell.text = text
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for r in p.runs:
            if bold:
                r.bold = True
            r.font.size = Pt(10)


def build_document(
    db,
    *,
    rnd: Round,
    country_label: str,
    promotion_limit: int,
    contest_id_filter: Optional[int],
    omit_email: bool,
) -> Document:
    variants = _country_variants(country_label)
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading(f"{country_label.title()} — Nomination report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph(f"Round: {rnd.name} (id={rnd.id})")
    doc.add_paragraph(
        f"Ranking rule: total vote stars (points), then shares, likes, comments, views, "
        f"then lower contestant id. Top {promotion_limit} in this country would migrate "
        f"on the next promotion step when the contest is on the country season "
        f"(country → regional)."
    )
    doc.add_paragraph("")

    contest_ids = [
        r[0]
        for r in db.execute(
            select(round_contests.c.contest_id).where(round_contests.c.round_id == rnd.id)
        ).fetchall()
    ]
    cq = db.query(Contest).filter(Contest.id.in_(contest_ids)).filter(Contest.contest_mode == "nomination")
    if contest_id_filter is not None:
        cq = cq.filter(Contest.id == contest_id_filter)
    contests = cq.order_by(Contest.id).all()

    for contest in contests:
        link = (
            db.query(ContestSeasonLink, ContestSeason)
            .join(ContestSeason, ContestSeason.id == ContestSeasonLink.season_id)
            .filter(
                ContestSeasonLink.contest_id == contest.id,
                ContestSeasonLink.is_active == True,
                ContestSeason.round_id == rnd.id,
            )
            .first()
        )
        if not link:
            continue
        _, season = link
        nxt = _next_level(season.level)

        doc.add_heading(f"{contest.name} (contest id {contest.id})", level=1)
        p_meta = doc.add_paragraph()
        p_meta.add_run(f"Active season level: {season.level.value} (season id {season.id}). ")
        if nxt:
            p_meta.add_run(f"Next migration step: → {nxt.value}. ")
        p_meta.add_run(f"Country filter (matched values): {sorted(variants)}.")

        ranked = _rank_tanzania_pool(
            db, season_id=season.id, contest_id=contest.id, country_variants=variants
        )
        if not ranked:
            doc.add_paragraph(f"No qualified contestants from {country_label} in this contest.")
            doc.add_paragraph("")
            continue

        n = len(ranked)
        doc.add_paragraph(f"Contestants in {country_label} for this contest: {n}")

        # Summary box
        top = ranked[:promotion_limit]
        top_names = [f"{i+1}. {_contestant_label(c)} ({stars}★)" for i, (c, _, stars) in enumerate(top)]
        p_sum = doc.add_paragraph()
        p_sum.add_run(f"Top {min(promotion_limit, n)} (would migrate): ").bold = True
        doc.add_paragraph(" ; ".join(top_names) if top_names else "(none)")

        if n > promotion_limit:
            rest = ranked[promotion_limit:]
            rest_line = ", ".join(_contestant_label(c) for c, _, _ in rest)
            doc.add_paragraph().add_run(f"Do not migrate (rank {promotion_limit + 1}–{n}): ").bold = True
            doc.add_paragraph(rest_line or "—")

        doc.add_paragraph("")

        # Table
        cols = [
            "Rank",
            "Contestant",
            "City",
            "Stars",
            "Shares",
            "Likes",
            "Comments",
            "Views",
            "Migrate",
        ]
        if not omit_email:
            cols.insert(3, "Email")

        table = doc.add_table(rows=1 + n, cols=len(cols))
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, h in enumerate(cols):
            _header_cell(hdr[i], h)

        for row_idx, (c, e, stars) in enumerate(ranked, start=1):
            row = table.rows[row_idx].cells
            rank = row_idx
            mig = "Yes" if rank <= promotion_limit else "No"
            col = 0
            _body_cell(row[col], str(rank), bold=(rank <= promotion_limit))
            col += 1
            _body_cell(row[col], _contestant_label(c), bold=(rank <= promotion_limit))
            col += 1
            if not omit_email:
                u = c.user
                em = (getattr(u, "email", None) or "") if u else ""
                _body_cell(row[col], em[:60])
                col += 1
            _body_cell(row[col], (c.city or "").strip()[:40])
            col += 1
            _body_cell(row[col], str(stars), bold=(rank <= promotion_limit))
            col += 1
            _body_cell(row[col], str(e.get("shares", 0)))
            col += 1
            _body_cell(row[col], str(e.get("likes", 0)))
            col += 1
            _body_cell(row[col], str(e.get("comments", 0)))
            col += 1
            _body_cell(row[col], str(e.get("views", 0)))
            col += 1
            _body_cell(row[col], mig, bold=(rank <= promotion_limit))

        doc.add_paragraph("")
        doc.add_paragraph("")

    doc.add_heading("Notes", level=2)
    doc.add_paragraph(
        "This report is a snapshot from the live database. It does not run migration. "
        "If the contest’s active season is not country, the same Tanzania filter still lists "
        "contestants whose profile country matches; migration rules apply when promotion runs "
        "from the country season."
    )
    return doc


def resolve_round(db, round_id: Optional[int], round_name: Optional[str]) -> Round:
    if round_id is not None:
        r = db.query(Round).filter(Round.id == round_id).first()
        if not r:
            sys.exit(f"Round id={round_id} not found")
        return r
    if round_name:
        r = db.query(Round).filter(Round.name.ilike(f"%{round_name.strip()}%")).first()
        if not r:
            sys.exit(f"No round matching name {round_name!r}")
        return r
    sys.exit("Provide --round-id or --round-name")


def main() -> None:
    ap = argparse.ArgumentParser(description="Word report: nomination contestants for one country")
    ap.add_argument("--round-id", type=int, default=None)
    ap.add_argument("--round-name", type=str, default=None)
    ap.add_argument("--country", type=str, default="Tanzania", help="Country name or short code (e.g. Tanzania, TZ)")
    ap.add_argument("--contest-id", type=int, default=None, help="Limit to one nomination contest")
    ap.add_argument("-o", "--output", required=True, help="Output .docx path")
    ap.add_argument("--promotion-limit", type=int, default=5, help="Top N count as migrating (default 5)")
    ap.add_argument("--no-email", action="store_true", help="Omit email column from tables")
    args = ap.parse_args()

    from app.db.session import SessionLocal

    db = SessionLocal()
    try:
        rnd = resolve_round(db, args.round_id, args.round_name)
        doc = build_document(
            db,
            rnd=rnd,
            country_label=args.country,
            promotion_limit=args.promotion_limit,
            contest_id_filter=args.contest_id,
            omit_email=args.no_email,
        )
        doc.save(args.output)
        print(f"Wrote Word report: {args.output!r}")
    finally:
        db.close()


if __name__ == "__main__":
    main()
